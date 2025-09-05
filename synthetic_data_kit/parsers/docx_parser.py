"""
Parser for Microsoft Word documents (DOCX).
"""
import os
import logging
from typing import Any, Optional

from synthetic_data_kit.parsers.base_parser import BaseParser

# Set up logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger("docx_parser")

class DOCXParser(BaseParser):
    """Parser for Microsoft Word documents (DOCX). Requires python-docx."""
    
    def parse(self, file_path: str) -> str:
        """
        Parse a DOCX file and return its content.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        logger.info(f"Parsing DOCX file: {file_path}")
        
        try:
            # Import docx module
            import docx
            
            # Open the document
            doc = docx.Document(file_path)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():
                                paragraphs.append(para.text)
            
            # Join paragraphs with double newlines
            content = "\n\n".join(paragraphs)
            
            # Clean the text
            content = self.clean_text(content)
            
            logger.info(f"Successfully parsed DOCX file ({len(content)} characters)")
            return content
        
        except ImportError:
            logger.error("DOCX parsing requires python-docx. Please install it.")
            raise ImportError("DOCX parsing requires python-docx to be installed.")
        
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {str(e)}")
            raise